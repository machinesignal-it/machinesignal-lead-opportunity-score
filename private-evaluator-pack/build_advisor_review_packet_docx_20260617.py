from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "private-evaluator-pack/MachineSignal_advisor_review_packet_20260617.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(31, 41, 55)
MUTED = RGBColor(90, 99, 113)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
PALE_GREEN = "EAF6EF"
PALE_RED = "FCECEC"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=None, size=9.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.1
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = color
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")
    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for w in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(w))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.tcW
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths[idx]))
            for margin_name in ("top", "bottom"):
                tc_mar = tc_pr.find(qn("w:tcMar"))
                if tc_mar is None:
                    tc_mar = OxmlElement("w:tcMar")
                    tc_pr.append(tc_mar)
                node = tc_mar.find(qn(f"w:{margin_name}"))
                if node is None:
                    node = OxmlElement(f"w:{margin_name}")
                    tc_mar.append(node)
                node.set(qn("w:w"), "80")
                node.set(qn("w:type"), "dxa")
            for margin_name in ("start", "end"):
                tc_mar = tc_pr.find(qn("w:tcMar"))
                node = tc_mar.find(qn(f"w:{margin_name}"))
                if node is None:
                    node = OxmlElement(f"w:{margin_name}")
                    tc_mar.append(node)
                node.set(qn("w:w"), "120")
                node.set(qn("w:type"), "dxa")


def add_rule(paragraph, color="D7DBE2", size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_para(doc, text="", style=None, bold=False, italic=False, color=None, size=None, after=6, before=0):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.1
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.name = "Calibri"
        if size:
            run.font.size = Pt(size)
        if color:
            run.font.color.rgb = color
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.1
        run = p.add_run(item)
        run.font.name = "Calibri"
        run.font.size = Pt(10.5)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.1
        run = p.add_run(item)
        run.font.name = "Calibri"
        run.font.size = Pt(10.5)


def add_status_table(doc, rows):
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    widths = [2500, 2200, 4660]
    set_table_geometry(table, widths)
    headers = ["Area", "Stato", "Significato"]
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, color=INK, size=9.5)
        set_cell_shading(table.rows[0].cells[idx], LIGHT_BLUE)
    for area, status, meaning in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], area, bold=True, size=9.3)
        set_cell_text(cells[1], status, size=9.3)
        set_cell_text(cells[2], meaning, size=9.3)
        if "Bloccato" in status or "No-go" in status:
            set_cell_shading(cells[1], PALE_RED)
        elif "Pronto" in status or "Consentita" in status:
            set_cell_shading(cells[1], PALE_GREEN)
    add_para(doc, "", after=4)


def add_question_section(doc, title, intro, groups):
    doc.add_heading(title, level=1)
    add_para(doc, intro, color=INK)
    for group_title, questions in groups:
        doc.add_heading(group_title, level=2)
        add_numbered(doc, questions)


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.text = "MachineSignal | Advisor Review Packet"
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = MUTED
    footer = section.footer.paragraphs[0]
    footer.text = "Documento preparatorio - nessuna attivazione commerciale"
    footer.runs[0].font.size = Pt(8.5)
    footer.runs[0].font.color.rgb = MUTED
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Masthead
    add_para(doc, "MACHINESIGNAL", bold=True, color=BLUE, size=10, after=2)
    title = add_para(doc, "Advisor Review Packet", bold=True, color=INK, size=24, after=3)
    subtitle = add_para(
        doc,
        "Riepilogo e domande per commercialista, legale e privacy",
        color=MUTED,
        size=13,
        after=12,
    )
    meta_rows = [
        ("Data", "17 giugno 2026"),
        ("Stato", "Sandbox/test - nessuna attivazione commerciale"),
        ("Scopo", "Preparare una eventuale beta controllata senza avviarla"),
        ("Cliente operativo", "Macchine: CRM, agenti AI, workflow, software e API client"),
    ]
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [1900, 7460])
    for label, value in meta_rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], label, bold=True, color=INK, size=9.5)
        set_cell_text(cells[1], value, size=9.5)
        set_cell_shading(cells[0], LIGHT_GRAY)
    add_para(doc, "", after=6)
    rule = add_para(doc, "", after=12)
    add_rule(rule)

    doc.add_heading("1. Sintesi", level=1)
    add_para(
        doc,
        "MachineSignal e' un progetto software/API pensato per essere usato soprattutto da sistemi automatici, non da persone che navigano un sito commerciale tradizionale. L'obiettivo e' permettere a CRM, agenti AI e workflow di scoprire target, valutare domini o aziende, ricevere score e preparare azioni successive in formato strutturato.",
        color=INK,
    )
    add_para(
        doc,
        "Il progetto e' tecnicamente pronto per lo scope sandbox attuale, ma non e' commercialmente live. Prima di incassare, fatturare o usare dati reali servono validazioni fiscali, legali e privacy.",
        bold=True,
        color=INK,
    )

    doc.add_heading("2. Stato Attuale", level=1)
    add_status_table(
        doc,
        [
            ("Sandbox tecnica", "Pronto per scope attuale", "API, documentazione, guardrail e status endpoint sono verificati."),
            ("Preparazione beta", "Consentita", "Gli agenti possono preparare materiali e policy senza attivare vendite."),
            ("Beta a pagamento", "No-go", "Mancano approvazioni fiscali, legali, pagamento, dati e supporto."),
            ("Go-live commerciale", "No-go", "Non e' pronto per vendita pubblica o marketplace."),
            ("Dati reali/personali", "Bloccato", "Nessuna lista cliente reale o dato personale nella fase attuale."),
            ("Outreach", "Bloccato", "Nessuna email o contatto esterno automatico."),
            ("Marketplace/MCP pubblico", "Bloccato", "Nessuna pubblicazione senza approvazione esplicita."),
        ],
    )

    doc.add_heading("3. Prodotti Ipotizzati", level=1)
    add_bullets(
        doc,
        [
            "Target Discovery Pack: ricerca di target aziendali coerenti con settore, area e criteri dati.",
            "Score Pack: valutazione di domini o aziende con punteggio e decisione operativa.",
            "Domain/Company Enrichment: arricchimento e classificazione di informazioni business.",
            "Deep Analysis: analisi piu' approfondita di opportunita', segnali e rischi.",
            "Action Pack: preparazione di azioni successive in formato strutturato per software o CRM.",
            "Opportunity Feed: scansioni periodiche e consegne ricorrenti.",
            "API Subscription: accesso mensile a limiti, crediti e funzioni API.",
        ],
    )

    doc.add_heading("4. Cosa Non E' Attivo", level=1)
    add_bullets(
        doc,
        [
            "Pagamenti reali.",
            "Fatture o ricevute.",
            "Raccolta di metodi di pagamento.",
            "API key di produzione.",
            "Liste clienti reali.",
            "Trattamento di dati personali.",
            "Campagne email o contatti commerciali esterni.",
            "Pubblicazione su marketplace, API directory, hosted MCP o registry pubblico.",
        ],
    )

    add_question_section(
        doc,
        "5. Domande Per Commercialista / Fisco",
        "Queste domande servono a capire cosa e' necessario prima di una eventuale beta a pagamento controllata.",
        [
            (
                "P.IVA e forma di partenza",
                [
                    "Possiamo fare una beta gratuita senza P.IVA se non incassiamo nulla?",
                    "Una beta a pagamento con pochi clienti richiede subito P.IVA o altra struttura?",
                    "Questa attivita' e' considerata continuativa fin dall'inizio?",
                    "Quale forma e' piu' adatta per vendere API/software online?",
                ],
            ),
            (
                "Fatturazione, IVA e ricavi",
                [
                    "Come si fatturano crediti, pacchetti API e abbonamenti mensili?",
                    "I crediti prepagati vanno fatturati subito o al consumo?",
                    "Che regole IVA valgono per Italia, UE ed extra UE?",
                    "Come gestire rimborsi o riaccredito crediti se l'output non e' valido?",
                ],
            ),
            (
                "Costi e adempimenti",
                [
                    "I costi di agenti AI, API esterne, Cloudflare, hosting e software sono deducibili?",
                    "Come vanno registrati crediti AI/API acquistati da fornitori esteri?",
                    "Quali documenti dobbiamo conservare per dimostrare costi e ricavi?",
                    "Cosa serve prima del primo incasso reale?",
                ],
            ),
        ],
    )

    add_question_section(
        doc,
        "6. Domande Per Legale / Privacy",
        "Queste domande servono a capire quali termini e regole dati servono prima di accettare clienti reali.",
        [
            (
                "Termini e responsabilita'",
                [
                    "Che termini servono per una API che fornisce score e analisi di opportunita'?",
                    "Come specifichiamo che lo score e' supporto decisionale e non garanzia di risultato?",
                    "Come limitiamo responsabilita' su errori, dati incompleti o decisioni prese dal cliente?",
                    "Serve un contratto diverso per beta gratuita e beta a pagamento?",
                ],
            ),
            (
                "Privacy, dati e DPA",
                [
                    "Se analizziamo solo domini aziendali pubblici, siamo comunque nel perimetro privacy?",
                    "Se il cliente carica una lista di aziende, quali obblighi abbiamo?",
                    "Se nella lista compaiono email, telefoni o nomi personali, cosa dobbiamo fare?",
                    "Serve una DPA/Data Processing Agreement per clienti business?",
                ],
            ),
            (
                "Macchine, sicurezza e comunicazione",
                [
                    "I termini devono prevedere che il cliente possa essere un software, agente AI o workflow automatico?",
                    "Come definiamo la responsabilita' se una macchina cliente compra crediti o richiama API automaticamente?",
                    "Che obblighi abbiamo se una API key viene esposta o riceviamo dati personali per errore?",
                    "Quali regole anti-spam e marketing B2B si applicano se in futuro volessimo fare comunicazione commerciale?",
                ],
            ),
        ],
    )

    doc.add_heading("7. Beta Controllata: Configurazione Minima", level=1)
    add_para(doc, "Se in futuro la beta a pagamento venisse autorizzata, la prima configurazione consigliata sarebbe molto prudente:", color=INK)
    add_bullets(
        doc,
        [
            "beta privata e controllata;",
            "approvazione manuale di ogni cliente;",
            "test sandbox prima di qualsiasi accesso produttivo;",
            "API key produzione emessa solo manualmente;",
            "limiti bassi di clienti, crediti e costo giornaliero;",
            "nessun dato personale;",
            "nessun outreach automatico;",
            "nessuna pubblicazione marketplace/MCP pubblico;",
            "kill switch e budget cap attivi;",
            "supporto con escalation limitata al proprietario.",
        ],
    )

    doc.add_heading("8. Risultato Atteso Dalla Consulenza", level=1)
    add_numbered(
        doc,
        [
            "Capire quale forma fiscale/amministrativa serve prima di vendere.",
            "Capire quali documenti fiscali servono prima del primo incasso.",
            "Capire quali termini, privacy policy o DPA servono prima di accettare clienti.",
            "Definire quali dati possiamo trattare e quali dobbiamo rifiutare.",
            "Definire le condizioni minime per una beta a pagamento sicura.",
        ],
    )

    doc.add_heading("9. Decisione Corrente", level=1)
    decision_rows = [
        ("Preparare materiali", "Si"),
        ("Chiedere parere consulenti", "Si"),
        ("Attivare beta a pagamento", "No"),
        ("Incassare denaro", "No"),
        ("Emettere fatture", "No"),
        ("Usare dati reali/personali", "No"),
        ("Contattare clienti", "No"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [4300, 5060])
    set_cell_text(table.rows[0].cells[0], "Decisione", bold=True, color=INK)
    set_cell_text(table.rows[0].cells[1], "Stato", bold=True, color=INK)
    set_cell_shading(table.rows[0].cells[0], LIGHT_BLUE)
    set_cell_shading(table.rows[0].cells[1], LIGHT_BLUE)
    for label, value in decision_rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], label, bold=True)
        set_cell_text(cells[1], value)
        set_cell_shading(cells[1], PALE_GREEN if value == "Si" else PALE_RED)

    doc.core_properties.author = "MachineSignal"
    doc.core_properties.title = "MachineSignal Advisor Review Packet"
    doc.core_properties.subject = "Riepilogo per consulenti fiscali, legali e privacy"
    doc.core_properties.comments = "Preparatory document. No commercial activation."

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
