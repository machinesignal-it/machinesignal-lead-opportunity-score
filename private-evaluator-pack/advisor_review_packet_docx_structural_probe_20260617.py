import json
from pathlib import Path
from zipfile import ZipFile

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "private-evaluator-pack" / "MachineSignal_advisor_review_packet_20260617.docx"
REPORT = ROOT / "private-evaluator-pack" / "advisor_review_packet_docx_structural_probe_report_20260617.md"
SUMMARY = ROOT / "private-evaluator-pack" / "advisor_review_packet_docx_structural_probe_summary_20260617.json"


def text_of_doc(doc):
    parts = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text.strip())
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text.strip())
    return "\n".join(parts)


def main():
    checks = []

    def add(name, passed, detail):
        checks.append({"name": name, "passed": bool(passed), "detail": detail})

    add("docx_exists", DOCX.exists(), str(DOCX))
    add("docx_non_empty", DOCX.exists() and DOCX.stat().st_size > 25000, str(DOCX.stat().st_size if DOCX.exists() else 0))

    doc = Document(DOCX)
    text = text_of_doc(doc)

    required_phrases = [
        "Advisor Review Packet",
        "Riepilogo e domande per commercialista, legale e privacy",
        "MachineSignal e' un progetto software/API",
        "Cosa Non E' Attivo",
        "Domande Per Commercialista / Fisco",
        "Domande Per Legale / Privacy",
        "Beta Controllata: Configurazione Minima",
        "Risultato Atteso Dalla Consulenza",
        "Attivare beta a pagamento",
        "Incassare denaro",
        "Usare dati reali/personali",
        "Contattare clienti",
    ]
    for phrase in required_phrases:
        add(f"contains_{phrase[:36].replace(' ', '_')}", phrase in text, phrase)

    required_no_values = [
        "Attivare beta a pagamento\nNo",
        "Incassare denaro\nNo",
        "Emettere fatture\nNo",
        "Usare dati reali/personali\nNo",
        "Contattare clienti\nNo",
    ]
    for phrase in required_no_values:
        add(f"decision_no_{phrase.splitlines()[0].replace(' ', '_')}", phrase in text, phrase)

    forbidden_phrases = [
        "Attivare beta a pagamento\nSi",
        "Incassare denaro\nSi",
        "Emettere fatture\nSi",
        "Usare dati reali/personali\nSi",
        "Contattare clienti\nSi",
        "beta a pagamento: si",
        "pagamenti reali attivi",
        "fatture attive",
    ]
    for phrase in forbidden_phrases:
        add(f"forbidden_absent_{phrase[:30].replace(' ', '_')}", phrase not in text, phrase)

    add("paragraph_count_reasonable", len(doc.paragraphs) >= 80, str(len(doc.paragraphs)))
    add("table_count_reasonable", len(doc.tables) >= 3, str(len(doc.tables)))
    add("section_count", len(doc.sections) == 1, str(len(doc.sections)))
    add("core_title", doc.core_properties.title == "MachineSignal Advisor Review Packet", str(doc.core_properties.title))
    add("core_author", doc.core_properties.author == "MachineSignal", str(doc.core_properties.author))

    with ZipFile(DOCX) as z:
        names = set(z.namelist())
        document_xml = z.read("word/document.xml").decode("utf-8")
        styles_xml = z.read("word/styles.xml").decode("utf-8")
    add("has_document_xml", "word/document.xml" in names, "word/document.xml")
    add("has_styles_xml", "word/styles.xml" in names, "word/styles.xml")
    add("has_table_grid", "TableGrid" in styles_xml or "Table Grid" in styles_xml, "Table Grid")
    add("has_heading_styles", "Heading1" in styles_xml and "Heading2" in styles_xml, "Heading1/Heading2")
    add("has_table_widths", 'w:tblW' in document_xml and 'w:gridCol' in document_xml, "tblW/gridCol")

    failed = [c for c in checks if not c["passed"]]
    status = "PASSED" if not failed else "FAILED"

    summary = {
        "date": "2026-06-17",
        "status": status,
        "checks": len(checks),
        "failed": len(failed),
        "docx": str(DOCX),
        "visual_render_qa": "not_completed_libreoffice_missing_word_com_hung",
    }
    SUMMARY.write_text(json.dumps(summary, indent=2), encoding="utf-8")

    report = "\n".join(
        [
            "# Advisor Review Packet DOCX - Structural Probe Report",
            "",
            "- Date: 2026-06-17",
            f"- Status: {status}",
            f"- Checks: {len(checks)}",
            f"- Failed: {len(failed)}",
            "- Visual render QA: not completed because LibreOffice/soffice is unavailable and Word COM export hung.",
            "",
            "## Result",
            "",
            "The DOCX exists, contains the required advisor packet sections, and keeps beta activation, money collection, invoices, real data and customer contact set to no.",
        ]
    )
    REPORT.write_text(report, encoding="utf-8")

    if failed:
        print(json.dumps(failed, indent=2))
        raise SystemExit(1)
    print(report)


if __name__ == "__main__":
    main()
